"""Shared synthetic-fixture factories for parsing unit tests.

These factories build small valid PDF and DOCX files from scratch so unit
tests never depend on the gitignored corpus — CI stays green with zero
external files.
"""

from pathlib import Path

import pytest


def _pdf_escape(text: str) -> str:
    """Escape characters that are special inside PDF literal strings."""
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def build_minimal_pdf_bytes(pages: list[str]) -> bytes:
    """Assemble a small valid single-xref PDF with one text line per page.

    Uses uncompressed content streams (``BT /F1 12 Tf (text) Tj ET`` with
    Helvetica) — pdfminer parses uncompressed streams without any filters.
    Object layout: 1=catalog, 2=pages tree, 3=font, then per page i:
    ``4+2i``=page object, ``5+2i``=content stream.
    """
    n = len(pages)
    objects: list[bytes] = []
    # obj 1: catalog
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    # obj 2: pages tree
    kids = " ".join(f"{4 + 2 * i} 0 R" for i in range(n))
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {n} >>".encode())
    # obj 3: font
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for i, text in enumerate(pages):
        content_num = 5 + 2 * i
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> "
                f"/Contents {content_num} 0 R >>"
            ).encode()
        )
        stream = f"BT /F1 12 Tf 72 720 Td ({_pdf_escape(text)}) Tj ET".encode()
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(out)
    count = len(objects) + 1
    out += f"xref\n0 {count}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    return bytes(out)


@pytest.fixture
def make_minimal_pdf():
    """Factory: write a small valid PDF with the given per-page text lines."""

    def _make(path: Path, pages: list[str]) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(build_minimal_pdf_bytes(pages))
        return path

    return _make


@pytest.fixture
def make_docx():
    """Factory: write a DOCX via python-docx.

    ``paragraphs`` entries are appended in order and may be:
    - ``str`` — a Normal-style paragraph
    - ``(text, style_name)`` tuple — a styled paragraph (e.g. "Heading 1")
    - ``{"table": rows}`` dict — a table with ``rows`` as list of list[str]

    ``tables`` (optional list of row-lists) are appended after all
    paragraph entries, as a convenience for non-interleaved cases.
    """

    def _make(path: Path, paragraphs: list, tables: list | None = None) -> Path:
        from docx import Document

        doc = Document()
        for item in paragraphs:
            if isinstance(item, dict) and "table" in item:
                _add_table(doc, item["table"])
            elif isinstance(item, tuple):
                text, style = item
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(item)
        for rows in tables or []:
            _add_table(doc, rows)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    def _add_table(doc, rows: list[list[str]]) -> None:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text

    return _make
